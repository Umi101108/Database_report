# -*- coding: utf-8 -*-
import datetime
from docx import Document
from docx.shared import Inches
import MySQLdb


class DatabaseDocument(object):

	def __init__(self):

		self.database = 'yaoxianzhi_2.1'
		try:
			self.conn = MySQLdb.connect(
				host='',
				port= 3306,
				user='root',
				passwd='',
				db='information_schema', 
				charset='utf8')
			self.cursor = self.conn.cursor()
		except MySQLdb.Error, e:
			print "连接数据库错误，原因%d: %s" % (e.args[0], e.args[1]) 


	def generateDocument(self, document_name="demo.docx"):

		document = Document()

		document.add_heading(u'数据库说明文档', 0)
		p = document.add_paragraph(u'表命名规范：bc-基础信息表， cr-用户相关表， gt-处方相关表， st-统计信息表')
		document.add_page_break()

		document.add_heading(u'表清单', 1)
		doc_table_list = document.add_table(rows=1, cols=2)
		doc_table_list.style = 'Table Grid'
		heading_cells = doc_table_list.rows[0].cells
		heading_cells[0].text, heading_cells[0].width = u'表名', Inches(3)
		heading_cells[1].text, heading_cells[1].width = u'注释', Inches(3)

		self.cursor.execute("select table_name,table_comment from information_schema.tables where table_schema='%s' and table_type='base table'" % self.database)
		tables = self.cursor.fetchall()
		for table in tables:
			cells = doc_table_list.add_row().cells
			cells[0].text, cells[1].text = table[0], table[1]

		for table in tables:
			doc_table_heading = document.add_heading('%s' % table[0], level=2)
			doc_table_heading.add_run('%s' % table[1]).bold = False
			print table[0], table[1]
			doc_table = document.add_table(rows=1, cols=6)
			doc_table.style = 'Table Grid'
			doc_table.allow_autofit = False
			heading_cells = doc_table.rows[0].cells
			heading_cells[0].text, heading_cells[0].width = u'名', Inches(6) * 4/15
			heading_cells[1].text, heading_cells[1].width = u'注释', Inches(6) * 5/15
			heading_cells[2].text, heading_cells[2].width = u'类型', Inches(6) * 3/15
			heading_cells[3].text, heading_cells[3].width = u'键', Inches(6) * 0.5/15
			heading_cells[4].text, heading_cells[4].width = u'可空', Inches(6) * 0.5/15
			heading_cells[5].text, heading_cells[5].width = u'默认', Inches(6) * 0.5/15

			self.cursor.execute("select COLUMN_NAME, COLUMN_COMMENT, COLUMN_TYPE, COLUMN_KEY, IS_NULLABLE, COLUMN_DEFAULT from information_schema.COLUMNS where table_schema='%s' and table_name='%s'" % (self.database, table[0]))
			table_columns = self.cursor.fetchall()
			for cols in table_columns:
				row_cells = doc_table.add_row().cells
				row_cells[0].text = cols[0]
				row_cells[1].text = cols[1]
				row_cells[2].text = cols[2]
				row_cells[3].text = cols[3]
				row_cells[4].text = cols[4]
				row_cells[5].text = str(cols[5]) if cols[5] else ''
			print "*"*20

		document.save(document_name + '_' + datetime.datetime.now().strftime('%Y%m%d') + '.docx')

	def main(self):
		document_name = 'yaoxianzhi_2.1'
		self.generateDocument(document_name)


if __name__ == "__main__":
	dd = DatabaseDocument()
	dd.main()